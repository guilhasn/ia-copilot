"""
Gera os 3 materiais da Sessão 1 em PDF com password.
Requer: python-docx, docx2pdf, PyMuPDF (fitz)
Uso:    python _gerar_materiais_s01.py
"""

import os
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

OUTPUT_DIR = Path(__file__).parent / "sessao-01"
PASSWORD = os.environ.get("ANFUP_S01_PWD")
if not PASSWORD:
    sys.exit("Defina a variável de ambiente ANFUP_S01_PWD antes de correr o script.")

NAVY = RGBColor(0x1F, 0x4E, 0x78)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
DARK = RGBColor(0x2C, 0x3E, 0x50)
RED = RGBColor(0xF4, 0x43, 0x36)
DARK_RED = RGBColor(0xB7, 0x1C, 0x1C)
AMBER = RGBColor(0xFF, 0x8F, 0x00)
GRAY = RGBColor(0x75, 0x75, 0x75)

HEADER_LEFT = "ANFUP — Associação Nacional dos Funcionários Universitários Portugueses"
HEADER_RIGHT = "Entidade Formadora Certificada DGERT"
FOOTER_LEFT = "IA — Aplicações ao trabalho das IES · Formador: Nuno Salvador"
FOOTER_RIGHT = "© 2026 Nuno Salvador"


def setup_styles(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = DARK
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.space_before = Pt(0)
    pf.line_spacing = 1.15


def add_header_footer(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    header = section.header
    hp = header.paragraphs[0]
    hp.text = HEADER_LEFT
    hp.style.font.size = Pt(8)
    hp.style.font.color.rgb = NAVY
    run = hp.add_run(f"\t{HEADER_RIGHT}")
    run.font.size = Pt(8)
    run.font.color.rgb = GREEN
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT

    hb = header.add_paragraph()
    hb.paragraph_format.space_before = Pt(2)
    hb.paragraph_format.space_after = Pt(6)
    border_run = hb.add_run()
    border_run.font.size = Pt(2)
    rPr = border_run._element
    pBdr = hb._element.get_or_add_pPr()
    bottom = pBdr.makeelement(qn("w:pBdr"), {})
    b = bottom.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "6",
            qn("w:space"): "1",
            qn("w:color"): "1F4E78",
        },
    )
    bottom.append(b)
    pBdr.append(bottom)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.text = FOOTER_LEFT
    fp.style.font.size = Pt(7.5)
    fp.style.font.color.rgb = GRAY
    run = fp.add_run(f"\t{FOOTER_RIGHT}")
    run.font.size = Pt(7.5)
    run.font.color.rgb = GRAY


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    run.font.size = Pt(22)
    run.font.color.rgb = NAVY
    run.bold = True

    if subtitle:
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(18)
        run2 = p2.add_run(subtitle)
        run2.font.size = Pt(12)
        run2.font.color.rgb = DARK
        run2.italic = True


def add_h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.size = Pt(14)
    run.font.color.rgb = NAVY
    run.bold = True


def add_h3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.size = Pt(11.5)
    run.font.color.rgb = NAVY
    run.bold = True


def add_para(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_quote(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    pBdr = p._element.get_or_add_pPr()
    border = pBdr.makeelement(qn("w:pBdr"), {})
    left = border.makeelement(
        qn("w:left"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "12",
            qn("w:space"): "8",
            qn("w:color"): "1F4E78",
        },
    )
    border.append(left)
    pBdr.append(border)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = NAVY
    return p


def set_cell_shading(cell, color_hex):
    shading = cell._element.get_or_add_tcPr()
    shd = shading.makeelement(
        qn("w:shd"),
        {qn("w:fill"): color_hex, qn("w:val"): "clear"},
    )
    shading.append(shd)


def to_pdf(docx_path):
    from docx2pdf import convert

    pdf_path = docx_path.with_suffix(".pdf")
    convert(str(docx_path), str(pdf_path))
    return pdf_path


def encrypt_pdf(pdf_path, password):
    import fitz

    doc = fitz.open(str(pdf_path))
    perm = fitz.PDF_PERM_PRINT | fitz.PDF_PERM_COPY | fitz.PDF_PERM_ANNOTATE
    tmp_path = pdf_path.with_suffix(".tmp.pdf")
    doc.save(
        str(tmp_path),
        encryption=fitz.PDF_ENCRYPT_AES_256,
        user_pw=password,
        owner_pw=password,
        permissions=perm,
    )
    doc.close()
    pdf_path.unlink()
    tmp_path.rename(pdf_path)


# ── DOCUMENTO 1: Matriz Semáforo ──────────────────────────────


def create_matriz():
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)

    add_title(doc, "Matriz Semáforo de Utilização de IA", "Guia rápido para staff das IES — o que posso pôr no Copilot e o que não")

    add_quote(doc, "Se hesitas a escrever isto num e-mail para uma pessoa fora da tua equipa → também não pões no Copilot.\nEm caso de dúvida → trata como AMARELO e pede orientação ao Encarregado de Proteção de Dados (EPD) da tua IES.")

    # Verde
    add_h2(doc, "🟢 VERDE — pode entrar à vontade")
    add_para(doc, "Conteúdo público, genérico ou já tratado. Usar Copilot livremente.")
    for item in [
        "Regulamentos, leis e normativos publicados",
        "Modelos genéricos de ofício, despacho ou parecer em branco",
        "Manuais de procedimentos internos sem dados pessoais",
        "Conteúdo já anonimizado (sem qualquer identificador direto ou indireto)",
        "Pedir ajuda técnica/conceptual sem dados reais",
        "Brainstorming sem referência a pessoas ou processos concretos",
        "Tradução de textos genéricos (ofícios-padrão, regulamentos publicados)",
    ]:
        add_bullet(doc, item)

    # Amarelo
    add_h2(doc, "🟡 AMARELO — entra, mas com salvaguardas")
    add_para(doc, "A maioria do trabalho real cabe aqui. Salvaguardas básicas: pseudonimizar, validar o output, não copiar a fundamentação direta.")
    for item in [
        "Trabalho preparatório (rascunhos de relatórios, draft de ofícios)",
        "Resumos de e-mails ou reuniões sobre temas operacionais sem conteúdo sensível",
        "Análise quantitativa agregada (totais, médias) sem identificação individual",
        "Documentos com referências indiretas a colegas",
        "Atas já aprovadas e divulgadas internamente",
        "Propostas de orçamento e prestação de contas (números agregados)",
        "Texto jurídico para reformular linguagem clara",
        "Apresentações executivas a partir de dados internos não confidenciais",
    ]:
        add_bullet(doc, item)

    # Vermelho
    add_h2(doc, "🔴 VERMELHO — não deves usar sem autorização do EPD")
    add_para(doc, "Toca direitos de pessoas ou tem efeito jurídico. Pede orientação antes.")
    for item in [
        "Dados pessoais identificáveis de candidatos, alunos ou trabalhadores em contexto de avaliação ou decisão",
        "Avaliações de júri em curso (concursos docentes, prémios, bolsas SAS, candidaturas)",
        "Propostas de menção SIADAP ou fundamentação de avaliação de desempenho",
        "Propostas concursais antes da decisão de adjudicação",
        "Cadernos de encargos antes da abertura pública",
        "Reclamações, queixas ou processos disciplinares em curso",
        "Documentos com etiqueta Purview \"Confidencial\" ou superior",
        "Dados de menores",
        "Conteúdo de procedimentos administrativos com prazo de defesa em curso",
    ]:
        add_bullet(doc, item)

    # Nunca
    add_h2(doc, "⛔ NUNCA — proibição absoluta")
    add_para(doc, "Independente de autorização, não pode entrar no Copilot.")
    for item in [
        "Dados especiais (RGPD art. 9.º): saúde, biometria, origem racial/étnica, opiniões políticas, convicções religiosas/filosóficas, vida ou orientação sexual, filiação sindical, dados genéticos",
        "Dados criminais (RGPD art. 10.º) — registo criminal, processos-crime, suspeitas",
        "Senhas, credenciais, chaves de API, tokens de sessão",
        "Conteúdo coberto por dever de sigilo profissional de terceiros (segredo médico, fiscal, advogado)",
        "Provas, enunciados de exame, critérios de correção antes da realização",
        "Conteúdo abrangido por NDA ou acordo de confidencialidade que exclua processamento por IA",
    ]:
        add_bullet(doc, item)

    doc.add_page_break()

    # CARA B
    add_h2(doc, "Como classifico o que estou prestes a colar?")
    add_para(doc, "1. O conteúdo identifica uma pessoa concreta (nome, NIF, n.º mecanográfico, e-mail)?")
    add_para(doc, "    → NÃO → VERDE ou AMARELO (depende de ser público ou interno)")
    add_para(doc, "    → SIM ↓")
    add_para(doc, "2. A IA pode influenciar uma decisão sobre essa pessoa (avaliação, classificação, admissão, sanção)?")
    add_para(doc, "    → SIM → VERMELHO ou NUNCA (verifica art. 9.º RGPD e art. 22.º RGPD)")
    add_para(doc, "    → NÃO ↓")
    add_para(doc, "3. O conteúdo cabe num e-mail interno entre a tua equipa?")
    add_para(doc, "    → SIM → AMARELO (pseudonimiza antes de colar)")
    add_para(doc, "    → NÃO → VERMELHO (pede ajuda ao EPD)")

    add_h2(doc, "Lembretes operacionais")
    for item in [
        "Usa sempre a tua conta institucional Microsoft 365. Nunca o Copilot/ChatGPT pessoal para trabalho.",
        "Desliga \"Pesquisa Web\" quando o prompt tiver conteúdo amarelo — força o Copilot a usar só o tenant.",
        "Verifica etiquetas Purview dos documentos referenciados — se vês \"Confidencial\", para.",
        "Valida o output — Copilot pode alucinar legislação. Confirma sempre referências legais.",
        "Pede sempre resposta em português europeu (\"responde em pt-pt\").",
        "Após terminar trabalho com conteúdo amarelo, apaga a conversa se não a precisares de manter.",
        "Nunca copies a fundamentação do Copilot diretamente para ato administrativo. A redação final é tua.",
    ]:
        add_bullet(doc, item)

    add_quote(doc, "A IA é assistente, não decisor.\nA responsabilidade pela decisão administrativa é, sempre, do trabalhador que assina o ato.")

    add_para(doc, "Versão 1.0 — junho 2026", italic=True, color=GRAY)

    path = OUTPUT_DIR / "Matriz_Semaforo.docx"
    doc.save(str(path))
    return path


# ── DOCUMENTO 2: Guia da Sessão 01 ─────────────────────────────


def create_guia():
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)

    add_title(doc, "Guia da Sessão 1", "Usar IA com critério no trabalho das IES")

    add_para(doc, "Formação: Inteligência Artificial — Aplicações ao trabalho das IES")
    add_para(doc, "Bloco 1 · Enquadramento e Literacia Crítica")
    add_para(doc, "Duração: 2 horas · Online síncrona")

    # Pilar 1
    add_h2(doc, "1. Literacia em IA — AI Act, artigo 4.º")

    add_para(doc, "O Regulamento (UE) 2024/1689 (AI Act) impõe, desde 2 de fevereiro de 2025, que todas as organizações garantam um nível suficiente de literacia em IA do seu pessoal.")

    add_para(doc, "A literacia em IA assenta em três eixos:")
    add_bullet(doc, "Conhecimento da tecnologia — o que é IA, RAG, alucinação, viés, limites do modelo.")
    add_bullet(doc, "Conhecimento do contexto de uso — que tarefas, sobre que dados, com que riscos para que pessoas.")
    add_bullet(doc, "Conhecimento da pessoa que utiliza — formação, experiência, função no fluxo de decisão.")

    add_para(doc, "A formação tem de ser proporcional ao risco. Não basta um vídeo de e-learning genérico — tem de existir formação específica ao contexto da organização e às ferramentas concretamente em uso.", italic=True)

    add_h3(doc, "O que a IES deve ter documentado")
    add_bullet(doc, "Política interna de uso de IA (assinada pelo dirigente máximo)")
    add_bullet(doc, "Plano de formação por função (mapeamento risco/perfil)")
    add_bullet(doc, "Registos de presença, conteúdos programáticos, materiais")
    add_bullet(doc, "Avaliação de conhecimentos")
    add_bullet(doc, "Atas de revisão anual da política")

    add_para(doc, "Para aprofundar: aiact-portugal.pt", italic=True, color=NAVY)

    # Pilar 2
    add_h2(doc, "2. Proteção de dados e decisões sobre pessoas — RGPD, artigo 22.º")

    add_quote(doc, "O titular dos dados tem o direito de não ficar sujeito a nenhuma decisão tomada exclusivamente com base no tratamento automatizado, incluindo a definição de perfis, que produza efeitos na sua esfera jurídica ou que o afete significativamente de forma similar.")

    add_h3(doc, "Apoio vs. decisão automatizada — quatro perguntas")
    add_bullet(doc, "A IA está apenas a ajudar a escrever ou está a influenciar a decisão?")
    add_bullet(doc, "O decisor consegue discordar do output?")
    add_bullet(doc, "O decisor conhece os dados e critérios usados?")
    add_bullet(doc, "Há tempo e condições reais para revisão humana?")

    add_h3(doc, "Aplicação prática nas IES")

    table = doc.add_table(rows=3, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Uso", "Avaliação"]
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
        set_cell_shading(cell, "1F4E78")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    data = [
        ("Copilot para organizar candidaturas, sumarizar CVs, identificar requisitos formais", "✅ Apoio (admissível com cuidados)"),
        ("Copilot para ranquear candidatos, propor classificações finais sem deliberação substantiva", "❌ Risco elevado de decisão automatizada"),
    ]
    for i, (col1, col2) in enumerate(data):
        table.rows[i + 1].cells[0].text = col1
        table.rows[i + 1].cells[1].text = col2

    # Pilar 3
    add_h2(doc, "3. Dever de sigilo e informação institucional — LTFP, artigo 73.º")

    add_para(doc, "O trabalhador em funções públicas tem dever de sigilo e dever de lealdade (Lei n.º 35/2014, art. 73.º/2). O incumprimento gera responsabilidade disciplinar.")

    add_h3(doc, "Copilot M365 — o que é preciso saber")
    add_bullet(doc, "Prompts, respostas e dados acedidos via Microsoft Graph não treinam foundation models.")
    add_bullet(doc, "Esta proteção aplica-se ao Copilot M365 com login organizacional — não ao Copilot pessoal.")
    add_bullet(doc, "A \"Pesquisa Web\" dentro do Copilot envia partes do prompt para fora do tenant — desativável.")
    add_bullet(doc, "Copilot vê o que o utilizador vê. Se há oversharing no SharePoint, o Copilot vai expor conteúdos mal partilhados.")
    add_bullet(doc, "Etiquetas Purview são respeitadas pelo Copilot e propagadas para o output.")

    add_quote(doc, "Antes do Copilot, o excesso de permissões podia ser um problema escondido. Com o Copilot, pode tornar-se pesquisável em linguagem natural.")

    # Checklist
    add_h2(doc, "E a minha IES? — 5 perguntas para levar ao EPD/DSI")

    add_bullet(doc, "1. A nossa IES tem uma política interna de utilização de IA aprovada e publicada?")
    add_bullet(doc, "2. As permissões do SharePoint e do OneDrive institucional foram revistas antes do roll-out do Copilot?")
    add_bullet(doc, "3. As etiquetas Purview estão configuradas e aplicadas aos documentos sensíveis?")
    add_bullet(doc, "4. Existe orientação clara sobre que processos não devem usar IA (SIADAP, júris, candidaturas)?")
    add_bullet(doc, "5. Há um procedimento para reportar incidentes ou utilizações inadequadas de IA?")

    # Referências
    add_h2(doc, "Referências")
    refs = [
        "Regulamento (UE) 2024/1689 — AI Act (EUR-Lex)",
        "AI Act em Portugal — aiact-portugal.pt",
        "RGPD — Regulamento (UE) 2016/679, artigo 22.º",
        "Lei n.º 35/2014 — LTFP, artigo 73.º",
        "Microsoft Learn — Enterprise data protection M365 Copilot",
        "Microsoft Learn — Purview for M365 Copilot",
    ]
    for ref in refs:
        add_bullet(doc, ref)

    path = OUTPUT_DIR / "Guia_Sessao_01.docx"
    doc.save(str(path))
    return path


# ── DOCUMENTO 3: Casos Práticos S01 ────────────────────────────


def create_casos():
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)

    add_title(doc, "Casos Práticos — Sessão 1", "Classifica cada cenário: Verde, Amarelo, Vermelho ou Nunca")

    add_para(doc, "Para cada cenário, classifica a utilização de IA segundo a Matriz Semáforo. A classificação depende dos dados usados, da fase do processo, da finalidade, do destinatário do output e do grau de influência da IA na decisão.", italic=True)

    casos = [
        {
            "n": 1,
            "titulo": "Ofício-padrão",
            "cenario": "Reformular um ofício-padrão de deferimento de matrícula, sem dados pessoais.",
            "cor": "Verde",
            "justificacao": "Trata-se de conteúdo genérico, sem identificação de pessoas e sem impacto decisório autónomo.",
        },
        {
            "n": 2,
            "titulo": "Legislação pública",
            "cenario": "Resumir a Lei Geral do Trabalho em Funções Públicas sobre deveres do trabalhador.",
            "cor": "Verde",
            "justificacao": "A legislação é pública. O uso da IA é de apoio à compreensão e síntese, exigindo sempre validação da fonte.",
        },
        {
            "n": 3,
            "titulo": "E-mails operacionais",
            "cenario": "Sumarizar 50 e-mails operacionais da semana para identificar assuntos pendentes.",
            "cor": "Amarelo",
            "justificacao": "Pode envolver informação interna, nomes, referências indiretas ou assuntos em curso. Deve haver minimização de dados e validação humana.",
        },
        {
            "n": 4,
            "titulo": "Ata de órgão universitário",
            "cenario": "Redigir rascunho de ata a partir de transcrição de reunião de órgão universitário.",
            "cor": "Amarelo",
            "justificacao": "Pode ser admissível como apoio à redação, mas exige cuidado com sigilo, nomes, deliberações, pontos reservados e revisão humana.",
        },
        {
            "n": 5,
            "titulo": "Propostas concursais",
            "cenario": "Analisar três propostas concursais e identificar a proposta de menor preço.",
            "cor": "Amarelo ou vermelho",
            "justificacao": "A classificação depende da fase do procedimento, da confidencialidade das propostas e do uso dado ao output. A IA pode apoiar extração e tabulação de critérios objetivos, mas não deve substituir a análise formal do júri.",
        },
        {
            "n": 6,
            "titulo": "Avaliação de desempenho",
            "cenario": "Gerar objetivos de avaliação de desempenho para uma trabalhadora identificada com base na autoavaliação.",
            "cor": "Vermelho",
            "justificacao": "Pode afetar uma pessoa identificada e influenciar uma decisão de avaliação. A IA pode apoiar formulações genéricas ou exemplos fictícios, mas não deve gerar fundamentação ou decisão sobre uma pessoa real.",
        },
        {
            "n": 7,
            "titulo": "Bolsa ou apoio social",
            "cenario": "Classificar candidatos a bolsa ou apoio social por ordem de carência.",
            "cor": "Vermelho",
            "justificacao": "Envolve decisão ou ordenação com impacto no acesso a benefício ou apoio. Pode envolver dados pessoais de elevada sensibilidade contextual. A IA não deve ordenar, selecionar ou excluir candidatos.",
        },
        {
            "n": 8,
            "titulo": "Atestado médico",
            "cenario": "Colar um atestado médico de um colega para pedir ao Copilot uma minuta de junta médica.",
            "cor": "Nunca",
            "justificacao": "Envolve dados de saúde identificáveis e informação altamente sensível. Este tipo de conteúdo não deve ser introduzido numa ferramenta de IA.",
        },
    ]

    for caso in casos:
        add_h3(doc, f"Caso {caso['n']} — {caso['titulo']}")
        add_para(doc, caso["cenario"])

        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        run = p.add_run("A minha classificação: ")
        run.font.size = Pt(10)
        run.font.color.rgb = GRAY
        run2 = p.add_run("______________________")
        run2.font.color.rgb = GRAY

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_before = Pt(8)
        run_label = p2.add_run(f"Resposta: {caso['cor']}")
        run_label.bold = True
        if "Verde" in caso["cor"]:
            run_label.font.color.rgb = GREEN
        elif "Amarelo" in caso["cor"] or "amarelo" in caso["cor"]:
            run_label.font.color.rgb = AMBER
        elif "Vermelho" in caso["cor"] or "vermelho" in caso["cor"]:
            run_label.font.color.rgb = RED
        elif "Nunca" in caso["cor"]:
            run_label.font.color.rgb = DARK_RED

        add_para(doc, caso["justificacao"], italic=True)

    add_para(doc, "")
    add_quote(doc, "A cor não depende apenas da tarefa. Depende dos dados usados, da fase do processo, da finalidade, do destinatário do output e do grau de influência da IA na decisão.")

    add_para(doc, "Versão 1.0 — junho 2026", italic=True, color=GRAY)

    path = OUTPUT_DIR / "Casos_Praticos_S01.docx"
    doc.save(str(path))
    return path


# ── MAIN ────────────────────────────────────────────────────────

if __name__ == "__main__":
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    print("A gerar Matriz Semáforo...")
    matriz_docx = create_matriz()
    print(f"  OK: {matriz_docx}")

    print("A gerar Guia da Sessão 01...")
    guia_docx = create_guia()
    print(f"  OK: {guia_docx}")

    print("A gerar Casos Práticos S01...")
    casos_docx = create_casos()
    print(f"  OK: {casos_docx}")

    print("\nA converter para PDF (requer Microsoft Word)...")
    pdfs = []
    for docx_path in [matriz_docx, guia_docx, casos_docx]:
        pdf = to_pdf(docx_path)
        print(f"  OK: {pdf}")
        pdfs.append(pdf)

    print(f"\nA proteger PDFs com password...")
    for pdf in pdfs:
        encrypt_pdf(pdf, PASSWORD)
        print(f"  LOCK: {pdf}")

    print("\nA limpar ficheiros .docx intermediários...")
    for docx_path in [matriz_docx, guia_docx, casos_docx]:
        docx_path.unlink()
        print(f"  DEL: {docx_path}")

    print("\nDONE!! Ficheiros em:", OUTPUT_DIR)
    for f in sorted(OUTPUT_DIR.glob("*.pdf")):
        size_kb = f.stat().st_size / 1024
        print(f"  - {f.name} ({size_kb:.0f} KB)")
