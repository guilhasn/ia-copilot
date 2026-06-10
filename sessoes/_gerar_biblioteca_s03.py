"""
Gera o ficheiro-biblioteca modelo da Sessão 3 (Biblioteca_de_Prompts.docx).
Documento editável, SEM password — é o template de trabalho dos formandos
(em especial dos que não têm licença Microsoft 365 Copilot).
Requer: python-docx
Uso:    python _gerar_biblioteca_s03.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT_DIR = Path(__file__).parent / "sessao-03"
OUTPUT_DIR.mkdir(exist_ok=True)

NAVY = RGBColor(0x1F, 0x4E, 0x78)
DARK = RGBColor(0x2C, 0x3E, 0x50)
GRAY = RGBColor(0x75, 0x75, 0x75)
LIGHT_BG = "D9E2F3"  # azul claro para células de rótulo

HEADER_LEFT = "ANFUP — Associação Nacional dos Funcionários Universitários Portugueses"
FOOTER_LEFT = "IA — Aplicações ao trabalho das IES · Formador: Nuno Salvador"
FOOTER_RIGHT = "© 2026 Nuno Salvador"

META_PROMPT = (
    "Atua como engenheiro de prompts para o Microsoft 365 Copilot numa "
    "instituição de ensino superior portuguesa.\n\n"
    "Vou descrever-te tarefas em linguagem corrente. Para cada tarefa:\n\n"
    "1. Reescreve-a como um prompt completo, com Objetivo, Contexto, Fonte "
    "e Expectativas.\n"
    "2. Antes de fechares o prompt, faz-me as perguntas necessárias sobre o "
    "que me faltou dizer — não inventes o que não sabes.\n"
    "3. Indica em que aplicação do Microsoft 365 devo usar o prompt, e em "
    "que modo (Trabalho ou Web).\n"
    "4. Termina com uma secção \"Validação:\" — o que devo verificar no "
    "resultado antes de o usar.\n\n"
    "Responde sempre em português de Portugal. Quando estiveres pronto, "
    "pede-me a primeira tarefa."
)

VERIFICADOR_PROMPT = (
    "Vou colar-te o resultado de uma pesquisa feita com IA. Verifica "
    "criticamente:\n\n"
    "1. As fontes citadas existem, ou precisam de confirmação manual?\n"
    "2. As datas são mesmo recentes, ou há relatórios antigos apresentados "
    "como atuais?\n"
    "3. A escala é proporcional, ou há generalizações a partir de casos "
    "isolados?\n"
    "4. Falta contraponto — vozes críticas, riscos, limitações?\n"
    "5. Que afirmações são genéricas e que afirmações são verificáveis "
    "(com números e fontes)?\n\n"
    "Termina com uma lista: \"Verificar manualmente antes de usar\".\n\n"
    "Resultado a verificar: [colar]"
)

CATEGORIAS = [
    ("Resumir", "e-mails, reuniões, documentos, cadeias de comunicação"),
    ("Redigir", "ofícios, respostas a e-mails, notas internas, pareceres"),
    ("Reformular", "melhorar clareza, ajustar tom, simplificar linguagem"),
    ("Analisar", "identificar pendências, comparar dados, extrair informação-chave"),
    ("Preparar", "agendas, pontos de situação, resumos executivos para reuniões"),
]


def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    tcPr.append(shd)


def setup_styles(doc):
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = DARK
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        header_p = section.header.paragraphs[0]
        header_p.text = HEADER_LEFT
        header_p.runs[0].font.size = Pt(8)
        header_p.runs[0].font.color.rgb = GRAY
        footer_p = section.footer.paragraphs[0]
        footer_p.text = f"{FOOTER_LEFT}\t\t{FOOTER_RIGHT}"
        footer_p.runs[0].font.size = Pt(8)
        footer_p.runs[0].font.color.rgb = GRAY


def add_title(doc, text, size=20, color=NAVY, center=True, bold=True):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    return p


def add_entry(doc, titulo, nome="", quando="", prompt="", validacao="",
              hint_nome="", hint_quando="", hint_prompt="", hint_validacao=""):
    h = doc.add_paragraph()
    run = h.add_run(titulo)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = NAVY

    table = doc.add_table(rows=4, cols=2)
    table.style = "Table Grid"
    table.columns[0].width = Cm(3.5)
    table.columns[1].width = Cm(13.5)

    rotulos = ["Nome", "Quando usar", "Prompt", "Validação"]
    valores = [nome, quando, prompt, validacao]
    hints = [hint_nome, hint_quando, hint_prompt, hint_validacao]

    for i, (rotulo, valor, hint) in enumerate(zip(rotulos, valores, hints)):
        c0 = table.cell(i, 0)
        c0.text = ""
        r = c0.paragraphs[0].add_run(rotulo)
        r.font.bold = True
        r.font.color.rgb = NAVY
        shade(c0, LIGHT_BG)

        c1 = table.cell(i, 1)
        c1.text = ""
        if valor:
            run = c1.paragraphs[0].add_run(valor)
            run.font.size = Pt(10)
        elif hint:
            run = c1.paragraphs[0].add_run(hint)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = GRAY

    doc.add_paragraph()


def main():
    doc = Document()
    setup_styles(doc)

    add_title(doc, "Biblioteca pessoal de prompts")
    add_title(doc, "IA — Aplicações ao trabalho das IES · Sessão 3",
              size=12, color=GRAY, bold=False)
    doc.add_paragraph()

    # Como usar
    h = doc.add_paragraph()
    r = h.add_run("Como usar este ficheiro")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = NAVY
    instrucoes = [
        "Este ficheiro é a sua biblioteca de prompts — não precisa de licença "
        "paga para o usar. Copie um prompt daqui e cole-o no Copilot Chat "
        "(m365.cloud.microsoft/chat, com a conta institucional) ou na versão "
        "web do Copilot (copilot.microsoft.com), direto no browser.",
        "Guarde o ficheiro no seu OneDrive. Para partilhar com a equipa, "
        "coloque uma cópia numa equipa do Teams ou no SharePoint do serviço — "
        "qualquer colega passa a poder copiar os prompts de lá.",
        "Cada entrada tem quatro campos: Nome (para encontrar depressa), "
        "Quando usar (a situação que o desencadeia), Prompt (o texto completo, "
        "pronto a colar) e Validação (o que verificar no resultado antes de o "
        "usar).",
        "A entrada #0 já vem preenchida: é o meta-prompt da Sessão 3 — o "
        "prompt que ajuda a criar todos os outros. Use-o para construir as "
        "entradas seguintes.",
        "IMPORTANTE — o Copilot Chat não memoriza instruções entre conversas: "
        "cada conversa nova começa do zero. Comece sempre por colar a entrada "
        "#0 (o meta-prompt) antes de descrever a tarefa. Sem este passo, as "
        "respostas voltam a ser genéricas.",
    ]
    for txt in instrucoes:
        p = doc.add_paragraph(txt)
        p.paragraph_format.space_after = Pt(6)
    doc.add_paragraph()

    # Entrada #0 — preenchida
    add_entry(
        doc,
        "Entrada #0 — O engenheiro de prompts (meta-prompt)",
        nome="Engenheiro de prompts (meta-prompt)",
        quando="Sempre que precisar de criar um prompt novo para a biblioteca, "
               "ou de melhorar um existente.",
        prompt=META_PROMPT,
        validacao="Rever o prompt gerado — cortar o que não se aplica, "
                  "confirmar aplicação e modo, testar uma vez antes de guardar.",
    )

    doc.add_page_break()

    # Entradas 1-5 — por categoria, com pistas
    for i, (cat, exemplos) in enumerate(CATEGORIAS, start=1):
        add_entry(
            doc,
            f"Entrada #{i} — Categoria: {cat}",
            hint_nome=f"Ex.: um nome curto que descreva a tarefa de {cat.lower()}",
            hint_quando=f"A situação concreta do seu serviço ({exemplos})",
            hint_prompt="Descreva a tarefa ao meta-prompt (entrada #0), responda "
                        "às perguntas dele e cole aqui o prompt final, já revisto "
                        "por si.",
            hint_validacao="O que vai verificar no resultado antes de o usar — "
                           "dados concretos, tom, completude.",
        )

    # Entrada #6 — verificador de pesquisas (preenchida)
    add_entry(
        doc,
        "Entrada #6 — O verificador de pesquisas",
        nome="Verificador de pesquisas com IA",
        quando="Depois de qualquer pesquisa feita com IA (Copilot em modo Web, "
               "ou outro assistente), antes de usar o resultado numa reunião, "
               "documento ou decisão.",
        prompt=VERIFICADOR_PROMPT,
        validacao="O verificador faz a triagem, mas a confirmação final das "
                  "fontes é sempre manual — ele diz onde olhar, não substitui "
                  "o olhar.",
    )

    # Entradas extra
    h = doc.add_paragraph()
    r = h.add_run("Entradas seguintes")
    r.font.size = Pt(13)
    r.font.bold = True
    r.font.color.rgb = NAVY
    doc.add_paragraph(
        "A biblioteca cresce com a prática: sempre que escrever um prompt que "
        "funcionou, acrescente-o aqui com os quatro campos. Antes de o partilhar "
        "com a equipa, faça-lhe o teste cruzado da Sessão 3: outra pessoa "
        "consegue usá-lo sem lhe fazer perguntas?"
    )

    out = OUTPUT_DIR / "Biblioteca_de_Prompts.docx"
    doc.save(out)
    print(f"OK: {out}")


if __name__ == "__main__":
    main()
